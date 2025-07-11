import os
import re
import datetime
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from config import TEMPLATE

load_dotenv()
api_key = os.getenv("OPEN_API_KEY")
if not api_key:
    raise ValueError("OPEN_API_KEY not set in .env")
client = OpenAI(api_key=api_key)


def generate_updated_resume(job_title: str, company_name: str, job_description: str):
    prompt = f"""
You are a professional resume enhancement assistant.

Given:
- Job Title: "{job_title}"
- Company Name: "{company_name}"
- Job Description: (see below)
- A sample summary and a structured tech‐stack format

Your tasks:
1. Extract all relevant keywords and skills from the job description.
2. Rewrite a professional resume summary (7–8 lines) that:
   - Mentions the job title and company name
   - Expresses interest in the position
   - Embeds the extracted keywords naturally
3. **Mandatory**: Update the Technical Skills section in this exact structure:
   Languages & Frameworks:
   Frontend:
   Backend:
   Cloud & DevOps:
   Databases:
   Tools & Other:
   Concepts:

⚠️ Format your response **exactly** like this:

Keywords: <comma-separated list>

Summary:
<7–8 line summary>

Tech Stack:
Languages & Frameworks: …
Frontend: …
Backend: …
Cloud & DevOps: …
Databases: …
Tools & Other: …
Concepts: …

Job Description:
\"\"\"
{job_description}
\"\"\"
"""
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    content = resp.choices[0].message.content

    print("RAW RESPONSE" + content)

    summary    = extract_summary(content)
    tech_stack = extract_tech_stack(content)
    return summary, tech_stack


def extract_summary(text: str) -> str:
    """Grab everything between 'Summary:' and 'Tech Stack:'."""
    m = re.search(r"Summary:\s*(.*?)\s*Tech Stack:", text, re.S)
    return m.group(1).strip() if m else ""


def extract_tech_stack(text: str) -> str:
    """Grab everything between 'Tech Stack:' and 'Job Description:'."""
    m = re.search(r"Tech Stack:\s*(.*?)\s*(Job Description:|$)", text, re.S)
    return m.group(1).strip() if m else ""


def _insert_multiline(paragraph, multiline: str):
    """Replace paragraph text with multiple runs, inserting a break between each line."""
    # clear existing text/runs
    paragraph.clear() if hasattr(paragraph, "clear") else setattr(paragraph, "text", "")
    lines = [line for line in multiline.splitlines() if line.strip()]
    if not lines:
        return
    # first line
    paragraph.add_run(lines[0])
    # subsequent lines with manual breaks
    for line in lines[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(line)


def update_existing_resume_docx(
    summary: str,
    tech_stack: str,
    job_title: str,
    company_name: str,
    template_path: str
) -> str:
    doc = Document(template_path)
    summary_found = tech_found = False

    for i, para in enumerate(doc.paragraphs):
        heading = para.text.strip().upper()

        # ── PROFESSIONAL SUMMARY ────────────────────────────────────
        if heading == "PROFESSIONAL SUMMARY":
            summary_found = True
            # clear until next Heading-style paragraph
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].style.name.startswith("Heading"):
                doc.paragraphs[j].text = ""
                j += 1
            # insert multi-line summary in first cleared paragraph
            _insert_multiline(doc.paragraphs[i + 1], summary)

        # ── TECHNICAL SKILLS ────────────────────────────────────────
        if heading == "TECHNICAL SKILLS":
            tech_found = True
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].style.name.startswith("Heading"):
                doc.paragraphs[j].text = ""
                j += 1
            _insert_multiline(doc.paragraphs[i + 1], tech_stack)

    if not summary_found:
        raise ValueError("Couldn't find 'PROFESSIONAL SUMMARY' in the template.")
    if not tech_found:
        raise ValueError("Couldn't find 'TECHNICAL SKILLS' in the template.")

    # ── Build output folder path: resume/YYYY-MM-DD ────────────────
    today = datetime.date.today().isoformat()
    out_dir = os.path.join("resume", today)
    os.makedirs(out_dir, exist_ok=True)

    company_name = sanitize_for_filename(company_name)
    job_title = sanitize_for_filename(job_title)

    # ── Save document into that folder ─────────────────────────────
    filename = f"Resume_harmeet_singh_{company_name}_{job_title}.docx"
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)

    print(f"✅ Resume updated and saved as: {out_path}")
    return out_path

def sanitize_for_filename(s: str) -> str:
    """
    Remove any non-printable characters and replace
    anything except letters, numbers, dot, dash or underscore with underscore.
    """
    # strip out zero-width and other non-printable chars
    s = ''.join(ch for ch in s if ch.isprintable())
    # replace disallowed chars
    return re.sub(r'[^A-Za-z0-9_\-\.]', '_', s)

# # ─── Example Usage ─────────────────────────────────────────────────────────────
# if __name__ == "__main__":
#     job_title    = "Software Developer Co-op"
#     company_name = "Rakuten Kobo"
#     job_description = """
# Rakuten Kobo Inc. is seeking a Software Developer Co-op to join our agile team.
# We are looking for students familiar with Python, REST APIs, GitHub, AWS, Docker, and CI/CD.
# Candidates should have experience in Agile environments, testing frameworks, and a passion for e-reading innovation.
# """
#     template = "Resume_Harmeet_Singh.docx"
#     summary, tech = generate_updated_resume(job_title, company_name, job_description)
#     update_existing_resume_docx(summary, tech, job_title, company_name, TEMPLATE)
